# -*- coding: utf-8 -*-
"""
spark_ingest_bronze.py (Bản đã vá lỗi + mở rộng theo yêu cầu 1-2-3-4)
------------------------------------------------------------
TẦNG BRONZE - OMNI-PARSER RICH SCHEMA

CÁC THAY ĐỔI SO VỚI BẢN GỐC:
  1. [FIX LỖI NGHIÊM TRỌNG] extract_ma_chi_tieu(): regex cũ chỉ nhận A-Z nên bỏ
     sót các mã có dấu tiếng Việt như 'ĐTMT01' -> mất toàn bộ chỉ tiêu Phòng Đào
     tạo một cách âm thầm. Đã đổi sang \\w Unicode để nhận đủ mọi mã.
  2. [Yêu cầu 1] parse_percent_or_number(): tách thêm cột numeric riêng cho
     muc_dang_ky / muc_dat, xử lý được cả số thập phân kiểu VN (dấu phẩy '3,5')
     lẫn kiểu quốc tế (dấu chấm '9.5'). Cột text gốc vẫn giữ nguyên để audit.
  3. [Yêu cầu 4] extract_quy_danh_gia_from_text(): trích kỳ đánh giá (Quý/Năm)
     từ NỘI DUNG văn bản thay vì hardcode "Q1/2026". Không phụ thuộc tên file,
     không giới hạn giá trị Quý/Năm cụ thể nào (nhận mọi Q1-Q4, mọi năm hợp lệ).
     Nếu không đọc được -> gán 'UNKNOWN_KY', KHÔNG chặn ingest, để lộ ra ở bước
     quality-gate của Silver (nessie_catalog_utils.check_quality_silver).
  4. extract_all_ky_candidates(): cảnh báo (không chặn) nếu 1 file có vẻ chứa
     lẫn nhiều kỳ khác nhau.
  5. checksum_sha256 giờ tính luôn cả quy_danh_gia để khóa duy nhất bản ghi
     phản ánh đúng khóa nghiệp vụ (ma_chi_tieu, quy_danh_gia).
------------------------------------------------------------
"""

import io
import os
import sys
import hashlib
import re
import zipfile
import boto3
from datetime import datetime
import pandas as pd
import pdfplumber
import docx

from env_config import MINIO_ENDPOINT, MINIO_ACCESS_KEY, MINIO_SECRET_KEY, MINIO_BUCKET_NAME

BUCKET_NAME    = MINIO_BUCKET_NAME
SOURCE_PREFIX  = "staging/"
ARCHIVE_PREFIX = "archive/"

KETQUA_KHONG_DAT     = "KHÔNG ĐẠT"
KETQUA_DAT           = "ĐẠT"
KETQUA_CHUA_DEN_KY   = "CHƯA ĐẾN KỲ ĐÁNH GIÁ"
QUY_DANH_GIA_UNKNOWN = "UNKNOWN_KY"


def get_s3_client():
    return boto3.client(
        "s3", endpoint_url=MINIO_ENDPOINT,
        aws_access_key_id=MINIO_ACCESS_KEY, aws_secret_access_key=MINIO_SECRET_KEY,
    )


def generate_checksum(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def clean_status_text(text: str) -> str:
    text = str(text).upper().strip()
    if re.search(r"(KHÔNG ĐẠT|FAILED)", text):
        return KETQUA_KHONG_DAT
    if re.search(r"(CHƯA ĐẾN KỲ|NOT DUE)", text):
        return KETQUA_CHUA_DEN_KY
    if re.search(r"(ĐẠT|PASSED|SUCCESS)", text):
        return KETQUA_DAT
    return text.replace("\n", " ").strip()


def extract_ma_chi_tieu(text: str):
    """Trích mã chỉ tiêu, ví dụ 'HT-MT05', 'ĐTMT01' -> 'ĐT-MT01'.

    [FIX] Dùng \\w với re.UNICODE (nhận mọi chữ cái Unicode, gồm Đ/Ă/Â/Ê/Ô/Ơ/Ư
    có dấu) thay vì [A-Z] như bản gốc — bản gốc bỏ sót hoàn toàn các mã như
    'ĐTMT01'..'ĐTMT13' (Phòng Đào tạo) vì chữ 'Đ' không thuộc A-Z.
    """
    text_clean = str(text).replace(" ", "").replace("\n", "").upper()
    match = re.search(r"([^\W\d_]{2,5}-?MT\d{2,3})", text_clean, re.UNICODE)
    if match:
        ma_raw = match.group(1)
        if "-" not in ma_raw:
            ma_raw = ma_raw.replace("MT", "-MT")
        return ma_raw
    return None


def parse_percent_or_number(text):
    """'100%' -> 100.0 | '3,5' -> 3.5 (số thập phân kiểu VN) | '9.5' -> 9.5
    | 'Chưa đến kỳ đánh giá' / 'ĐẠT' -> None (không phải giá trị số)."""
    if text is None:
        return None
    cleaned = str(text).strip().replace("%", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def extract_quy_danh_gia_from_text(full_text: str):
    """Trích kỳ đánh giá từ NỘI DUNG văn bản (không phụ thuộc tên file).
    Nhận MỌI Quý 1-4 và MỌI năm hợp lệ (không hardcode 1 giá trị cụ thể nào),
    ví dụ: 'QUÝ 2/2022' -> 'Q2/2022', 'QUÝ 3/2014' -> 'Q3/2014'.
    Chỉ cần pattern xuất hiện ở BẤT KỲ đâu trong toàn bộ file (gộp mọi trang),
    không bắt buộc lặp lại ở từng trang.
    """
    text_upper = str(full_text).upper()

    m = re.search(r"QUÝ\s*(\d)\s*/\s*(\d{4})", text_upper)
    if m:
        quy, nam = int(m.group(1)), int(m.group(2))
        if 1 <= quy <= 4 and 2000 <= nam <= 2100:
            return f"Q{quy}/{nam}"

    m2 = re.search(r"\bQ(\d)\s*QUÝ\b", text_upper)
    year_m = re.search(r"NĂM\s*(\d{4})", text_upper)
    if m2 and year_m:
        quy, nam = int(m2.group(1)), int(year_m.group(1))
        if 1 <= quy <= 4 and 2000 <= nam <= 2100:
            return f"Q{quy}/{nam}"

    return None


def extract_all_ky_candidates(full_text: str) -> set:
    """Tìm TẤT CẢ các kỳ khác nhau xuất hiện trong file, để cảnh báo (KHÔNG
    chặn ingest) nếu 1 file có vẻ chứa lẫn nhiều kỳ đánh giá khác nhau."""
    text_upper = str(full_text).upper()
    found = set()
    for match in re.finditer(r"QUÝ\s*(\d)\s*/\s*(\d{4})", text_upper):
        quy, nam = int(match.group(1)), int(match.group(2))
        if 1 <= quy <= 4 and 2000 <= nam <= 2100:
            found.add(f"Q{quy}/{nam}")
    return found


def get_full_text_pdf(pdf) -> str:
    return "\n".join(page.extract_text() or "" for page in pdf.pages)


def get_full_text_docx(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def parse_pdf(file_bytes: bytes):
    """Trả về (rows, quy_danh_gia, ky_candidates)."""
    rows = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        full_text = get_full_text_pdf(pdf)
        quy_danh_gia = extract_quy_danh_gia_from_text(full_text)
        ky_candidates = extract_all_ky_candidates(full_text)

        for page in pdf.pages:
            table = page.extract_table()
            if table:
                for r in table:
                    if not r or len(r) < 8:
                        continue
                    ma = extract_ma_chi_tieu(r[0])
                    if ma:
                        noi_dung = str(r[1]).strip().replace("\n", " ") if r[1] else "N/A"
                        dinh_ky = str(r[2]).strip().replace("\n", " ") if r[2] else "N/A"
                        muc_dang_ky = str(r[3]).strip().replace("\n", " ") if r[3] else "N/A"
                        muc_dat = str(r[4]).strip().replace("\n", " ") if r[4] else "N/A"
                        ket_qua = clean_status_text(r[5])
                        nguyen_nhan = str(r[6]).strip().replace("\n", " ") if r[6] else ""
                        hanh_dong = str(r[7]).strip().replace("\n", " ") if r[7] else ""
                        rows.append((ma, noi_dung, dinh_ky, muc_dang_ky, muc_dat, ket_qua, nguyen_nhan, hanh_dong))
    return rows, quy_danh_gia, ky_candidates


def parse_docx(file_bytes: bytes):
    """Trả về (rows, quy_danh_gia, ky_candidates)."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except (zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise ValueError("Invalid or corrupt DOCX file") from exc

    full_text = get_full_text_docx(doc)
    quy_danh_gia = extract_quy_danh_gia_from_text(full_text)
    ky_candidates = extract_all_ky_candidates(full_text)

    rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 8:
                continue
            ma = extract_ma_chi_tieu(cells[0])
            if ma:
                noi_dung = cells[1].replace("\n", " ") if cells[1] else "N/A"
                dinh_ky = cells[2].replace("\n", " ") if cells[2] else "N/A"
                muc_dang_ky = cells[3].replace("\n", " ") if cells[3] else "N/A"
                muc_dat = cells[4].replace("\n", " ") if cells[4] else "N/A"
                ket_qua = clean_status_text(cells[5])
                nguyen_nhan = cells[6].replace("\n", " ") if cells[6] else ""
                hanh_dong = cells[7].replace("\n", " ") if cells[7] else ""
                rows.append((ma, noi_dung, dinh_ky, muc_dang_ky, muc_dat, ket_qua, nguyen_nhan, hanh_dong))
    return rows, quy_danh_gia, ky_candidates


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    s3_client = get_s3_client()
    extracted_data = []
    response = s3_client.list_objects_v2(Bucket=BUCKET_NAME, Prefix=SOURCE_PREFIX)
    if "Contents" not in response:
        sys.exit(0)

    processed_keys = []
    for obj in response["Contents"]:
        file_key = obj["Key"]

        if file_key.endswith("/"):
            continue
        ext = os.path.splitext(file_key)[1].lower()
        file_bytes = s3_client.get_object(Bucket=BUCKET_NAME, Key=file_key)["Body"].read()

        raw_rows, quy_danh_gia, ky_candidates = [], None, set()
        if ext == ".pdf":
            raw_rows, quy_danh_gia, ky_candidates = parse_pdf(file_bytes)
        elif ext == ".docx":
            try:
                raw_rows, quy_danh_gia, ky_candidates = parse_docx(file_bytes)
            except ValueError as exc:
                print(f"WARNING: Bỏ qua file không hợp lệ {file_key}: {exc}")
        else:
            print(f"SKIP: Định dạng không hỗ trợ cho file {file_key}")

        if len(ky_candidates) > 1:
            print(
                f"⚠️  CẢNH BÁO: file '{file_key}' có vẻ chứa nhiều kỳ khác nhau {ky_candidates}, "
                f"chỉ đang gán '{quy_danh_gia}' cho toàn bộ dữ liệu trong file này. "
                f"Cần kiểm tra thủ công."
            )

        quy_danh_gia_final = quy_danh_gia or QUY_DANH_GIA_UNKNOWN
        if quy_danh_gia_final == QUY_DANH_GIA_UNKNOWN and raw_rows:
            print(
                f"⚠️  CẢNH BÁO: không xác định được kỳ đánh giá trong file '{file_key}'. "
                f"Dữ liệu vẫn được ingest vào Bronze với quy_danh_gia='{QUY_DANH_GIA_UNKNOWN}' "
                f"để admin kiểm tra thủ công ở bước Silver (KHÔNG chặn upload)."
            )

        for ma, noi_dung, dk, m_dk, m_dat, kq, nguyen_nhan, hanh_dong in raw_rows:
            nhom = ma.split("-")[0]
            checksum = generate_checksum(
                f"{file_key}_{ma}_{quy_danh_gia_final}_{dk}_{m_dk}_{m_dat}_{kq}"
            )
            extracted_data.append({
                "file_nguon": os.path.basename(file_key),
                "ma_chi_tieu": ma,
                "nhom_don_vi": nhom,
                "quy_danh_gia": quy_danh_gia_final,
                "noi_dung_muc_tieu": noi_dung,
                "dinh_ky_thu_thap": dk,
                "muc_dang_ky": m_dk,
                "muc_dang_ky_numeric": parse_percent_or_number(m_dk),
                "muc_dat": m_dat,
                "muc_dat_numeric": parse_percent_or_number(m_dat),
                "ket_qua_he_thong": kq,
                "nguyen_nhan": nguyen_nhan,
                "hanh_dong_khac_phuc": hanh_dong,
                "checksum_sha256": checksum,
            })

        processed_keys.append(file_key)

    if extracted_data:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_key = f"bronze/data_extracted_{timestamp}.parquet"

        df = pd.DataFrame(extracted_data)
        parquet_buffer = io.BytesIO()
        df.to_parquet(parquet_buffer, index=False, engine="pyarrow")
        s3_client.put_object(Bucket=BUCKET_NAME, Key=output_key, Body=parquet_buffer.getvalue())
        print(f"Đã tạo Parquet: {output_key} ({len(extracted_data)} dòng)")
    else:
        print("Không có dữ liệu hợp lệ nào được trích xuất để ghi Parquet.")

    if processed_keys:
        for key in processed_keys:
            archive_key = key.replace(SOURCE_PREFIX, ARCHIVE_PREFIX, 1)
            s3_client.copy_object(
                Bucket=BUCKET_NAME,
                CopySource={"Bucket": BUCKET_NAME, "Key": key},
                Key=archive_key,
            )
            s3_client.delete_object(Bucket=BUCKET_NAME, Key=key)

        print(f"HOÀN THÀNH INGEST! Đã dọn dẹp {len(processed_keys)} files khỏi staging.")
    else:
        print("Không có file nào mới để xử lý.")


if __name__ == "__main__":
    main()